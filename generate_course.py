"""
Script de génération du support de cours TravelDock
Génère un document Word (.docx) complet en français
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Styles globaux ──────────────────────────────────────────────────────────

def set_page_margins(doc, top=2, bottom=2, left=2.5, right=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)

set_page_margins(doc)

# Style Normal
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

# Heading 1
h1 = doc.styles['Heading 1']
h1.font.name = 'Calibri'
h1.font.size = Pt(20)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x15, 0x63, 0xC7)

# Heading 2
h2 = doc.styles['Heading 2']
h2.font.name = 'Calibri'
h2.font.size = Pt(16)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

# Heading 3
h3 = doc.styles['Heading 3']
h3.font.name = 'Calibri'
h3.font.size = Pt(13)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x1E, 0x88, 0xE5)

# ── Helpers ─────────────────────────────────────────────────────────────────

def h(level, text):
    doc.add_heading(text, level=level)

def p(text=''):
    para = doc.add_paragraph(text)
    para.style = doc.styles['Normal']
    return para

def code_block(text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(6)
    # fond gris via shading
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F0F0F0')
    pPr.append(shd)
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    return para

def note(text, color='FFF9C4'):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Cm(0.5)
    para.paragraph_format.right_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color)
    pPr.append(shd)
    run = para.add_run('💡 ' + text)
    run.font.size = Pt(10)
    run.font.italic = True
    return para

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def numbered(text):
    doc.add_paragraph(text, style='List Number')

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE DE TITRE
# ══════════════════════════════════════════════════════════════════════════════

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Cm(4)
run = title_para.add_run('TravelDock')
run.font.size = Pt(40)
run.font.bold = True
run.font.color.rgb = RGBColor(0x15, 0x63, 0xC7)

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle_para.add_run('Application de Gestion de Voyages')
run2.font.size = Pt(22)
run2.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

doc.add_paragraph()

tech_para = doc.add_paragraph()
tech_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = tech_para.add_run('Symfony 7 · Docker · Doctrine ORM · MySQL 8')
run3.font.size = Pt(14)
run3.font.italic = True
run3.font.color.rgb = RGBColor(0x1E, 0x88, 0xE5)

doc.add_paragraph()
doc.add_paragraph()

info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = info_para.add_run(f'Support de cours — {datetime.date.today().strftime("%B %Y")}')
run4.font.size = Pt(12)
run4.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SOMMAIRE (manuel)
# ══════════════════════════════════════════════════════════════════════════════

h(1, 'Sommaire')
sommaire = [
    ('1', 'Introduction au projet TravelDock'),
    ('2', 'Docker — Containerisation de l\'environnement'),
    ('  2.1', 'Qu\'est-ce que Docker ?'),
    ('  2.2', 'Les containers du projet'),
    ('  2.3', 'Le Dockerfile PHP 8.3'),
    ('  2.4', 'La configuration Nginx'),
    ('  2.5', 'Le fichier docker-compose.yml'),
    ('  2.6', 'Commandes Docker essentielles'),
    ('3', 'Symfony 7 — Framework PHP moderne'),
    ('  3.1', 'Présentation de Symfony'),
    ('  3.2', 'Structure du projet'),
    ('  3.3', 'Le Kernel et le point d\'entrée'),
    ('  3.4', 'Les fichiers de configuration'),
    ('  3.5', 'Les variables d\'environnement (.env)'),
    ('4', 'Doctrine ORM — Mapping objet-relationnel'),
    ('  4.1', 'Principe de l\'ORM'),
    ('  4.2', 'Configuration de Doctrine'),
    ('  4.3', 'L\'Entity Manager'),
    ('5', 'Les Entités — Modèle de données'),
    ('  5.1', 'L\'entité Destination'),
    ('  5.2', 'L\'entité Travel'),
    ('  5.3', 'L\'entité Booking'),
    ('  5.4', 'Les relations entre entités'),
    ('  5.5', 'Les contraintes de validation'),
    ('6', 'Les Repositories — Requêtes personnalisées'),
    ('  6.1', 'Principe des Repositories'),
    ('  6.2', 'Le QueryBuilder'),
    ('  6.3', 'Exemples de requêtes'),
    ('7', 'Les Migrations — Versionnement du schéma'),
    ('  7.1', 'Principe des migrations'),
    ('  7.2', 'Créer et exécuter une migration'),
    ('  7.3', 'Analyse de la migration TravelDock'),
    ('8', 'Les Controllers et les Routes'),
    ('  8.1', 'Architecture MVC dans Symfony'),
    ('  8.2', 'Le routing par attributs PHP'),
    ('  8.3', 'Le HomeController'),
    ('  8.4', 'Le TravelController (CRUD)'),
    ('  8.5', 'Le DestinationController'),
    ('  8.6', 'Le BookingController'),
    ('9', 'Les Templates Twig'),
    ('  9.1', 'Le moteur de templates Twig'),
    ('  9.2', 'Le template de base'),
    ('  9.3', 'Héritage de templates'),
    ('10', 'Mise en pratique — Exercices'),
    ('11', 'Référence des commandes'),
]
for num, title in sommaire:
    para = doc.add_paragraph()
    run = para.add_run(f'{num}  {title}')
    if not num.startswith('  '):
        run.font.bold = True
        run.font.size = Pt(11)
    else:
        run.font.size = Pt(10)
        para.paragraph_format.left_indent = Cm(1)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 1 — INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

h(1, '1. Introduction au projet TravelDock')

p("""TravelDock est une application web de gestion de voyages développée avec Symfony 7.
Elle permet de gérer des destinations, des voyages et des réservations.
Ce support de cours vous guide, étape par étape, dans la création de cette application
en partant de zéro, avec Docker comme environnement de développement.""")

p()
h(2, 'Objectifs pédagogiques')
bullet('Comprendre la containerisation avec Docker')
bullet('Maîtriser les fondamentaux de Symfony 7')
bullet('Utiliser Doctrine ORM pour interagir avec une base de données MySQL')
bullet('Créer des entités, des migrations et des controllers')
bullet('Construire des templates Twig')
bullet('Implémenter un CRUD complet (Create, Read, Update, Delete)')

p()
h(2, 'Architecture technique')
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Technologie'
hdr[1].text = 'Version'
hdr[2].text = 'Rôle'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.font.bold = True

data = [
    ('PHP', '8.3', 'Langage serveur'),
    ('Symfony', '7.1', 'Framework PHP'),
    ('Doctrine ORM', '3.x', 'Accès base de données'),
    ('MySQL', '8.0', 'Base de données relationnelle'),
    ('Nginx', '1.25', 'Serveur web'),
    ('Docker', '28+', 'Containerisation'),
    ('Twig', '3.x', 'Moteur de templates'),
    ('Bootstrap', '5.3', 'CSS framework'),
]
for row_data in data:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 2 — DOCKER
# ══════════════════════════════════════════════════════════════════════════════

h(1, '2. Docker — Containerisation de l\'environnement')

h(2, '2.1 Qu\'est-ce que Docker ?')
p("""Docker est une plateforme open-source qui permet d'empaqueter une application et
toutes ses dépendances dans des unités standardisées appelées containers.
Chaque container est isolé, léger, et portable : il fonctionnera de la même façon
sur n'importe quelle machine, qu'il s'agisse du PC d'un développeur, d'un serveur
de test ou de production.""")

p()
note("Un container Docker ≠ une machine virtuelle. Les containers partagent le noyau de l'OS hôte, ce qui les rend beaucoup plus légers et rapides à démarrer.")

p()
h(3, 'Concepts clés')
bullet('Image : template immuable qui définit le contenu d\'un container (OS, logiciels, config)')
bullet('Container : instance en cours d\'exécution d\'une image')
bullet('Dockerfile : fichier texte décrivant comment construire une image')
bullet('docker-compose.yml : fichier décrivant et orchestrant plusieurs containers')
bullet('Volume : mécanisme de persistance des données entre redémarrages')
bullet('Réseau : communication entre containers')

p()
h(2, '2.2 Les containers du projet TravelDock')
p('Notre projet utilise 4 containers orchestrés avec Docker Compose :')

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
for i, txt in enumerate(['Container', 'Image', 'Port', 'Rôle']):
    hdr2[i].text = txt
    for run in hdr2[i].paragraphs[0].runs:
        run.font.bold = True

data2 = [
    ('symfony_php',        'php:8.3-fpm (custom)',  '9000 (interne)',  'Exécute le code PHP'),
    ('symfony_nginx',      'nginx:1.25-alpine',      '8080 → 80',       'Serveur web / reverse proxy'),
    ('symfony_mysql',      'mysql:8.0',              '3306 → 3306',     'Base de données MySQL'),
    ('symfony_phpmyadmin', 'phpmyadmin:5.2',         '8081 → 80',       'Interface web MySQL'),
]
for row_data in data2:
    row = table2.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

p()
h(2, '2.3 Le Dockerfile PHP 8.3')
p("""Le Dockerfile est le fichier qui décrit comment construire notre image PHP
personnalisée avec toutes les extensions requises par Symfony :""")

code_block("""FROM php:8.3-fpm

# Installation des dépendances système
RUN apt-get update && apt-get install -y \\
    git curl libpng-dev libonig-dev libxml2-dev \\
    libzip-dev libicu-dev zip unzip \\
    && docker-php-ext-install \\
        pdo pdo_mysql mbstring exif \\
        pcntl bcmath gd zip intl opcache \\
    && apt-get clean

# Copie de Composer depuis l'image officielle
COPY --from=composer:latest /usr/bin/composer /usr/bin/composer

# Configuration PHP (mémoire, upload, opcache)
RUN echo "memory_limit=512M" >> /usr/local/etc/php/conf.d/custom.ini \\
    && echo "upload_max_filesize=64M" >> /usr/local/etc/php/conf.d/custom.ini

WORKDIR /var/www/symfony

# Utilisateur non-root pour la sécurité
RUN groupadd -g 1000 www && useradd -u 1000 -ms /bin/bash -g www www
USER www

EXPOSE 9000""")

p()
h(3, 'Explication ligne par ligne')

table3 = doc.add_table(rows=1, cols=2)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Instruction'
hdr3[1].text = 'Signification'
for cell in hdr3:
    for run in cell.paragraphs[0].runs:
        run.font.bold = True

data3 = [
    ('FROM php:8.3-fpm', 'Image de base officielle PHP 8.3 avec PHP-FPM'),
    ('RUN apt-get install', 'Installation des bibliothèques système nécessaires'),
    ('docker-php-ext-install', 'Compilation et activation des extensions PHP'),
    ('pdo_mysql', 'Pilote PDO pour MySQL (requis par Doctrine)'),
    ('intl', 'Extension internationalisation (requis par Symfony)'),
    ('COPY --from=composer', 'Multi-stage build : copie Composer depuis son image officielle'),
    ('WORKDIR', 'Définit le dossier de travail par défaut dans le container'),
    ('USER www', 'Lance PHP-FPM avec un utilisateur non-root (sécurité)'),
    ('EXPOSE 9000', 'Documente le port PHP-FPM (utilisé par Nginx)'),
]
for row_data in data3:
    row = table3.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]

p()
h(2, '2.4 La configuration Nginx')
p("""Nginx joue le rôle de serveur web et de reverse proxy. Il reçoit les requêtes HTTP
sur le port 80 et les transmet à PHP-FPM via le protocole FastCGI sur le port 9000 :""")

code_block("""server {
    listen 80;
    root /var/www/symfony/public;   # Dossier public de Symfony

    location / {
        # Toutes les URLs passent par index.php (Front Controller)
        try_files $uri /index.php$is_args$args;
    }

    location ~ ^/index\\.php(/|$) {
        fastcgi_pass php:9000;       # Transmission à PHP-FPM
        fastcgi_split_path_info ^(.+\\.php)(/.*)$;
        include fastcgi_params;
        fastcgi_param SCRIPT_FILENAME $realpath_root$fastcgi_script_name;
        internal;
    }

    # Sécurité : refus direct des .php autres qu'index.php
    location ~ \\.php$ {
        return 404;
    }
}""")

note("Le pattern try_files $uri /index.php$is_args$args est fondamental dans Symfony : il permet au Front Controller (index.php) de gérer toutes les URLs, même celles qui ne correspondent pas à un fichier physique.")

p()
h(2, '2.5 Le fichier docker-compose.yml')
p("""docker-compose.yml orchestre tous les containers et définit leurs interactions.
Voici le fichier complet annoté :""")

code_block("""services:
  php:
    build:
      context: ./docker/php     # Chemin vers le Dockerfile
      dockerfile: Dockerfile
    container_name: symfony_php
    volumes:
      - ./symfony:/var/www/symfony  # Montage du code source
    depends_on:
      mysql:
        condition: service_healthy  # Attend que MySQL soit prêt

  nginx:
    image: nginx:1.25-alpine
    ports:
      - "8080:80"               # Port hôte:port container
    volumes:
      - ./symfony:/var/www/symfony
      - ./docker/nginx/default.conf:/etc/nginx/conf.d/default.conf
    depends_on:
      - php

  mysql:
    image: mysql:8.0
    environment:
      MYSQL_ROOT_PASSWORD: rootpassword
      MYSQL_DATABASE: travelDock
      MYSQL_USER: symfony
      MYSQL_PASSWORD: symfony_password
    ports:
      - "3306:3306"
    volumes:
      - mysql_data:/var/lib/mysql  # Persistance des données
    healthcheck:
      test: ["CMD", "mysqladmin", "ping", "-h", "localhost"]
      interval: 10s
      retries: 5

  phpmyadmin:
    image: phpmyadmin:5.2
    environment:
      PMA_HOST: mysql
    ports:
      - "8081:80"

volumes:
  mysql_data:           # Volume nommé pour MySQL""")

p()
h(2, '2.6 Commandes Docker essentielles')

table4 = doc.add_table(rows=1, cols=2)
table4.style = 'Table Grid'
hdr4 = table4.rows[0].cells
hdr4[0].text = 'Commande'
hdr4[1].text = 'Description'
for cell in hdr4:
    for run in cell.paragraphs[0].runs:
        run.font.bold = True

cmds = [
    ('docker compose up -d',              'Démarre tous les containers en arrière-plan'),
    ('docker compose up -d --build',      'Reconstruit les images et démarre'),
    ('docker compose down',               'Arrête et supprime les containers'),
    ('docker compose ps',                 'Liste les containers et leur statut'),
    ('docker compose logs -f',            'Affiche les logs en temps réel'),
    ('docker compose exec php bash',      'Ouvre un terminal dans le container PHP'),
    ('docker compose exec php composer install', 'Lance Composer dans le container PHP'),
    ('docker compose build --no-cache',   'Reconstruit sans utiliser le cache'),
]
for row_data in cmds:
    row = table4.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 3 — SYMFONY 7
# ══════════════════════════════════════════════════════════════════════════════

h(1, '3. Symfony 7 — Framework PHP moderne')

h(2, '3.1 Présentation de Symfony')
p("""Symfony est un framework PHP open-source créé par SensioLabs en 2005.
La version 7 (LTS) apporte des performances accrues, le support natif des attributs PHP 8
pour le routing et la validation, et une architecture encore plus modulaire basée sur
des composants réutilisables.""")

p()
h(3, 'Philosophie de Symfony')
bullet('Convention over Configuration : des conventions intelligentes réduisent la configuration')
bullet('DRY (Don\'t Repeat Yourself) : factorisation et réutilisabilité du code')
bullet('Composants découplés : chaque composant peut être utilisé indépendamment')
bullet('Autowiring : injection automatique des dépendances par type-hinting')
bullet('Tests intégrés : architecture pensée pour la testabilité')

p()
h(2, '3.2 Structure du projet')
code_block("""symfony-docker/
├── docker/                    ← Configuration Docker
│   ├── php/Dockerfile
│   ├── nginx/default.conf
│   └── mysql/init.sql
├── symfony/                   ← Code source Symfony
│   ├── bin/
│   │   └── console            ← CLI Symfony
│   ├── config/
│   │   ├── packages/          ← Config des bundles (doctrine.yaml, framework.yaml…)
│   │   ├── routes.yaml        ← Configuration du routing
│   │   └── services.yaml      ← Définition des services
│   ├── migrations/            ← Migrations de base de données
│   ├── public/
│   │   └── index.php          ← Front Controller (point d'entrée unique)
│   ├── src/
│   │   ├── Controller/        ← Controllers HTTP
│   │   ├── Entity/            ← Entités Doctrine
│   │   ├── Repository/        ← Repositories Doctrine
│   │   └── Kernel.php         ← Noyau de l'application
│   ├── templates/             ← Templates Twig
│   ├── var/
│   │   ├── cache/             ← Cache compilé
│   │   └── log/               ← Journaux d'application
│   ├── vendor/                ← Dépendances Composer
│   ├── .env                   ← Variables d'environnement
│   └── composer.json          ← Dépendances PHP
└── docker-compose.yml""")

p()
h(2, '3.3 Le Kernel et le point d\'entrée')
p("""Le Kernel est le cœur de Symfony. Il initialise le container de services,
charge la configuration et orchestre le cycle de vie des requêtes HTTP.""")

code_block("""<?php
// src/Kernel.php
namespace App;

use Symfony\\Bundle\\FrameworkBundle\\Kernel\\MicroKernelTrait;
use Symfony\\Component\\HttpKernel\\Kernel as BaseKernel;

class Kernel extends BaseKernel
{
    use MicroKernelTrait;  // Charge config/, routes.yaml, services.yaml automatiquement
}""")

p()
p('Le Front Controller est le fichier public/index.php, point d\'entrée unique de toutes les requêtes :')
code_block("""<?php
// public/index.php
use App\\Kernel;

require_once dirname(__DIR__).'/vendor/autoload_runtime.php';

return function (array $context) {
    return new Kernel($context['APP_ENV'], (bool) $context['APP_DEBUG']);
};""")

note("Le Front Controller est le seul fichier PHP accessible directement depuis le web. Nginx redirige toutes les requêtes vers ce fichier.")

p()
h(2, '3.4 Les fichiers de configuration')
p('Symfony 7 utilise des fichiers YAML pour configurer ses composants. Voici les principaux :')

bullet('config/packages/framework.yaml : configuration du noyau (sessions, CSRF…)')
bullet('config/packages/doctrine.yaml : configuration de l\'ORM Doctrine')
bullet('config/packages/twig.yaml : configuration du moteur de templates')
bullet('config/routes.yaml : découverte automatique des routes')
bullet('config/services.yaml : définition et autowiring des services')

p()
h(2, '3.5 Les variables d\'environnement (.env)')
p("""Le fichier .env contient les variables de configuration sensibles ou spécifiques
à l'environnement. Il ne doit JAMAIS être commité avec des secrets en production.""")

code_block("""# .env
APP_ENV=dev                 # Environnement (dev, prod, test)
APP_SECRET=a1b2c3d4e5f6    # Clé secrète (CSRF, cookies...)

# URL de connexion à la base de données
DATABASE_URL="mysql://symfony:symfony_password@mysql:3306/travelDock?serverVersion=8.0&charset=utf8mb4"
#                      user       password      host  port  database""")

note("En production, les variables d'environnement sont souvent injectées par le serveur (Docker secrets, variables système) et non lues depuis .env.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 4 — DOCTRINE ORM
# ══════════════════════════════════════════════════════════════════════════════

h(1, '4. Doctrine ORM — Mapping objet-relationnel')

h(2, '4.1 Principe de l\'ORM')
p("""Un ORM (Object-Relational Mapping) est une technique de programmation qui permet
de manipuler une base de données relationnelle à travers des objets PHP,
sans écrire de SQL manuellement.

Doctrine ORM fait le lien entre :
  • les classes PHP (Entity) ↔ les tables MySQL
  • les propriétés PHP ↔ les colonnes
  • les instances d'objets ↔ les lignes""")

p()
h(3, 'Avantages de Doctrine')
bullet('Abstraction : le même code PHP fonctionne avec MySQL, PostgreSQL, SQLite…')
bullet('Typage fort : les données sont des objets PHP, pas des tableaux')
bullet('Lazy loading : les relations sont chargées à la demande')
bullet('Cache de requêtes : optimisation automatique des performances')
bullet('Migrations : versionnement du schéma de base de données')

p()
h(2, '4.2 Configuration de Doctrine')
p('La configuration Doctrine dans config/packages/doctrine.yaml définit comment accéder à la base :')

code_block("""doctrine:
    dbal:
        url: '%env(resolve:DATABASE_URL)%'  # Lit la variable DATABASE_URL du .env
    orm:
        auto_generate_proxy_classes: true
        naming_strategy: doctrine.orm.naming_strategy.underscore_number_aware
        auto_mapping: true
        mappings:
            App:
                type: attribute        # Utilise les attributs PHP #[ORM\\...] dans les entités
                dir: '%kernel.project_dir%/src/Entity'
                prefix: 'App\\Entity'""")

p()
h(2, '4.3 L\'Entity Manager')
p("""L'EntityManagerInterface est le service principal de Doctrine. Il gère le cycle de vie
des entités : création, modification, suppression et requêtes.""")

code_block("""// Injection par autowiring dans un Controller
public function __construct(
    private readonly EntityManagerInterface $entityManager,
) {}

// Persister (préparer l'insertion)
$destination = new Destination();
$destination->setName('Tokyo');
$this->entityManager->persist($destination);

// Flush : exécute le SQL (INSERT/UPDATE/DELETE)
$this->entityManager->flush();

// Supprimer
$this->entityManager->remove($destination);
$this->entityManager->flush();""")

note("persist() met l'entité dans la \"file d'attente\". flush() exécute réellement le SQL. Cette distinction permet de grouper plusieurs opérations dans une seule transaction.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 5 — LES ENTITÉS
# ══════════════════════════════════════════════════════════════════════════════

h(1, '5. Les Entités — Modèle de données')

p("""Une entité Symfony est une classe PHP annotée avec des attributs Doctrine (#[ORM\\...]).
Elle représente une table dans la base de données.
TravelDock possède 3 entités : Destination, Travel et Booking.""")

p()
h(2, '5.1 L\'entité Destination')
p("""La Destination représente un lieu de voyage (ville/pays).
C'est l'entité la plus simple, elle est le point de départ de la relation.""")

code_block("""<?php
namespace App\\Entity;

use App\\Repository\\DestinationRepository;
use Doctrine\\Common\\Collections\\ArrayCollection;
use Doctrine\\Common\\Collections\\Collection;
use Doctrine\\ORM\\Mapping as ORM;
use Symfony\\Component\\Validator\\Constraints as Assert;

#[ORM\\Entity(repositoryClass: DestinationRepository::class)]
#[ORM\\Table(name: 'destination')]
class Destination
{
    #[ORM\\Id]
    #[ORM\\GeneratedValue]
    #[ORM\\Column]
    private ?int $id = null;               // Clé primaire auto-incrémentée

    #[ORM\\Column(length: 100)]
    #[Assert\\NotBlank]                    // Validation : champ obligatoire
    private ?string $name = null;          // Colonne VARCHAR(100)

    #[ORM\\Column(length: 100)]
    private ?string $country = null;

    #[ORM\\Column(type: 'text', nullable: true)]
    private ?string $description = null;   // Colonne TEXT, nullable

    #[ORM\\Column(length: 10, nullable: true)]
    private ?string $iataCode = null;      // Code aéroport (ex: CDG, NRT)

    // Relation One-to-Many : une Destination a plusieurs Travels
    #[ORM\\OneToMany(targetEntity: Travel::class, mappedBy: 'destination')]
    private Collection $travels;

    public function __construct()
    {
        $this->travels = new ArrayCollection();  // Initialisation obligatoire
    }

    // Getters / Setters...
    public function __toString(): string
    {
        return $this->name . ' (' . $this->country . ')';
    }
}""")

p()
h(2, '5.2 L\'entité Travel')
p('Travel représente un voyage proposé, avec ses dates, son prix et sa destination.')

code_block("""#[ORM\\Entity(repositoryClass: TravelRepository::class)]
class Travel
{
    #[ORM\\Id]
    #[ORM\\GeneratedValue]
    #[ORM\\Column]
    private ?int $id = null;

    #[ORM\\Column(length: 150)]
    #[Assert\\NotBlank]
    #[Assert\\Length(min: 3, max: 150)]
    private ?string $title = null;

    // DATE_MUTABLE = colonne SQL DATE (sans heure)
    #[ORM\\Column(type: Types::DATE_MUTABLE)]
    private ?\\DateTimeInterface $departureDate = null;

    #[ORM\\Column(type: Types::DATE_MUTABLE)]
    // Validation cross-champ : returnDate > departureDate
    #[Assert\\GreaterThan(propertyPath: 'departureDate')]
    private ?\\DateTimeInterface $returnDate = null;

    // DECIMAL(10,2) pour les prix (évite les erreurs de flottants)
    #[ORM\\Column(type: Types::DECIMAL, precision: 10, scale: 2)]
    #[Assert\\Positive]
    private ?string $price = null;

    #[ORM\\Column]
    private ?int $maxParticipants = null;

    #[ORM\\Column(length: 20)]
    private string $status = 'draft';      // Valeur par défaut

    // Relation Many-to-One : un Travel appartient à une Destination
    #[ORM\\ManyToOne(targetEntity: Destination::class, inversedBy: 'travels')]
    #[ORM\\JoinColumn(nullable: false)]    // FK obligatoire
    private ?Destination $destination = null;

    // Relation One-to-Many avec cascade : supprime les réservations si le voyage est supprimé
    #[ORM\\OneToMany(targetEntity: Booking::class, mappedBy: 'travel', cascade: ['persist', 'remove'])]
    private Collection $bookings;

    // Méthode métier : calcul de la durée
    public function getDuration(): int
    {
        return $this->departureDate->diff($this->returnDate)->days;
    }

    // Méthode métier : places disponibles
    public function getAvailableSeats(): int
    {
        $booked = 0;
        foreach ($this->bookings as $booking) {
            if ($booking->getStatus() !== 'cancelled') {
                $booked += $booking->getNumberOfPersons();
            }
        }
        return $this->maxParticipants - $booked;
    }
}""")

p()
h(2, '5.3 L\'entité Booking')
p('Booking représente une réservation effectuée par un voyageur pour un voyage donné.')

code_block("""#[ORM\\Entity(repositoryClass: BookingRepository::class)]
class Booking
{
    // Constantes de statut pour éviter les \"magic strings\"
    public const STATUS_PENDING   = 'pending';
    public const STATUS_CONFIRMED = 'confirmed';
    public const STATUS_CANCELLED = 'cancelled';

    #[ORM\\Id]
    #[ORM\\GeneratedValue]
    #[ORM\\Column]
    private ?int $id = null;

    #[ORM\\Column(length: 100)]
    #[Assert\\NotBlank]
    private ?string $firstName = null;

    #[ORM\\Column(length: 180)]
    #[Assert\\Email]                       // Validation format email
    private ?string $email = null;

    #[ORM\\Column]
    #[Assert\\LessThanOrEqual(value: 10)]  // Max 10 personnes par réservation
    private ?int $numberOfPersons = null;

    #[ORM\\Column(type: Types::DECIMAL, precision: 10, scale: 2)]
    private ?string $totalPrice = null;

    #[ORM\\Column(length: 20)]
    private string $status = self::STATUS_PENDING;  // Référence à la constante

    // Many-to-One : une Booking appartient à un Travel
    #[ORM\\ManyToOne(targetEntity: Travel::class, inversedBy: 'bookings')]
    #[ORM\\JoinColumn(nullable: false)]
    private ?Travel $travel = null;

    public function __construct()
    {
        $this->bookedAt = new \\DateTime();  // Date automatique à la création
    }

    // Calcul automatique du prix total
    public function calculateTotalPrice(): void
    {
        if ($this->travel && $this->numberOfPersons) {
            $this->totalPrice = (string) ((float) $this->travel->getPrice()
                                          * $this->numberOfPersons);
        }
    }
}""")

p()
h(2, '5.4 Les relations entre entités')
p('TravelDock utilise deux types de relations Doctrine :')

h(3, 'ManyToOne / OneToMany (relation bidirectionnelle)')
p("""La relation la plus courante. Un Travel appartient à UNE Destination (ManyToOne).
Une Destination peut avoir PLUSIEURS Travels (OneToMany).""")

code_block("""// Côté "Many" : Travel a une clé étrangère vers Destination
#[ORM\\ManyToOne(targetEntity: Destination::class, inversedBy: 'travels')]
#[ORM\\JoinColumn(nullable: false)]
private ?Destination $destination = null;

// Côté "One" : Destination liste ses Travels (pas de colonne SQL)
#[ORM\\OneToMany(targetEntity: Travel::class, mappedBy: 'destination')]
private Collection $travels;""")

note("La propriété mappedBy='destination' et inversedBy='travels' forment les deux faces d'une même relation. La colonne FK (destination_id) n'existe que du côté ManyToOne (Travel).")

p()
h(2, '5.5 Les contraintes de validation')
p("""Symfony Validator permet de définir des règles de validation directement
sur les entités via des attributs PHP :""")

code_block("""use Symfony\\Component\\Validator\\Constraints as Assert;

// Champ obligatoire, non vide
#[Assert\\NotBlank(message: 'Le titre est obligatoire.')]
private ?string $title = null;

// Longueur min/max
#[Assert\\Length(min: 3, max: 150, minMessage: 'Minimum 3 caractères')]
private ?string $title = null;

// Validation email
#[Assert\\Email(message: 'Email invalide')]
private ?string $email = null;

// Valeur positive
#[Assert\\Positive(message: 'Le prix doit être positif')]
private ?string $price = null;

// Comparaison entre champs
#[Assert\\GreaterThan(propertyPath: 'departureDate')]
private ?\\DateTimeInterface $returnDate = null;

// Valeur max
#[Assert\\LessThanOrEqual(value: 10)]
private ?int $numberOfPersons = null;""")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 6 — LES REPOSITORIES
# ══════════════════════════════════════════════════════════════════════════════

h(1, '6. Les Repositories — Requêtes personnalisées')

h(2, '6.1 Principe des Repositories')
p("""Un Repository est une classe dédiée aux requêtes de base de données pour une entité donnée.
Il étend ServiceEntityRepository et hérite de méthodes de base comme findAll(), find(),
findBy(), findOneBy(). On y ajoute ensuite des méthodes de requêtes métier personnalisées.""")

code_block("""<?php
namespace App\\Repository;

use App\\Entity\\Destination;
use Doctrine\\Bundle\\DoctrineBundle\\Repository\\ServiceEntityRepository;
use Doctrine\\Persistence\\ManagerRegistry;

class DestinationRepository extends ServiceEntityRepository
{
    public function __construct(ManagerRegistry $registry)
    {
        parent::__construct($registry, Destination::class);
    }
    // Méthodes héritées disponibles :
    // find(id), findAll(), findBy(['country' => 'France']), findOneBy([...])
}""")

p()
h(2, '6.2 Le QueryBuilder')
p("""Le QueryBuilder est l'outil principal pour construire des requêtes complexes
en PHP (sans écrire de SQL brut) :""")

code_block("""// DestinationRepository.php

// Recherche par pays (égalité exacte)
public function findByCountry(string $country): array
{
    return $this->createQueryBuilder('d')      // 'd' = alias de Destination
        ->andWhere('d.country = :country')     // Condition WHERE
        ->setParameter('country', $country)    // Liaison du paramètre (anti-injection SQL)
        ->orderBy('d.name', 'ASC')             // Tri
        ->getQuery()                           // Génère l'objet Query
        ->getResult();                         // Exécute et retourne un tableau d'objets
}

// Recherche partielle (LIKE)
public function search(string $query): array
{
    return $this->createQueryBuilder('d')
        ->andWhere('d.name LIKE :query OR d.country LIKE :query')
        ->setParameter('query', '%' . $query . '%')   // % = wildcard SQL
        ->orderBy('d.name', 'ASC')
        ->getQuery()
        ->getResult();
}

// Jointure : destinations avec voyages publiés
public function findWithActiveTravel(): array
{
    return $this->createQueryBuilder('d')
        ->innerJoin('d.travels', 't')          // JOIN sur la relation travels
        ->andWhere('t.status = :status')
        ->setParameter('status', 'published')
        ->groupBy('d.id')
        ->getQuery()
        ->getResult();
}""")

p()
h(2, '6.3 Exemples du TravelRepository')

code_block("""// TravelRepository.php

// Eager loading : charge les relations en une seule requête (évite les N+1 queries)
public function findWithBookings(): array
{
    return $this->createQueryBuilder('t')
        ->leftJoin('t.bookings', 'b')
        ->addSelect('b')                       // Inclut les bookings dans la requête
        ->leftJoin('t.destination', 'd')
        ->addSelect('d')
        ->orderBy('t.departureDate', 'ASC')
        ->getQuery()
        ->getResult();
}

// Requête avec date dynamique (voyages futurs uniquement)
public function findPublished(): array
{
    return $this->createQueryBuilder('t')
        ->andWhere('t.status = :status')
        ->setParameter('status', 'published')
        ->andWhere('t.departureDate > :now')
        ->setParameter('now', new \\DateTime())  // Date actuelle
        ->orderBy('t.departureDate', 'ASC')
        ->leftJoin('t.destination', 'd')
        ->addSelect('d')
        ->getQuery()
        ->getResult();
}

// BookingRepository — agrégation (SUM)
public function countPersonsForTravel(int $travelId): int
{
    $result = $this->createQueryBuilder('b')
        ->select('SUM(b.numberOfPersons)')       // Agrégation
        ->andWhere('b.travel = :travelId')
        ->setParameter('travelId', $travelId)
        ->andWhere('b.status != :status')
        ->setParameter('status', 'cancelled')
        ->getQuery()
        ->getSingleScalarResult();               // Retourne une seule valeur
    return (int) $result;
}""")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 7 — LES MIGRATIONS
# ══════════════════════════════════════════════════════════════════════════════

h(1, '7. Les Migrations — Versionnement du schéma')

h(2, '7.1 Principe des migrations')
p("""Les migrations sont des fichiers PHP qui décrivent de manière incrémentale
les modifications apportées au schéma de la base de données.
Elles permettent de versionner la structure de la BDD avec le code, et d'appliquer
les changements de façon reproductible sur n'importe quel environnement.""")

bullet('Chaque migration a un numéro de version unique (timestamp)')
bullet('La méthode up() applique les changements')
bullet('La méthode down() permet d\'annuler les changements (rollback)')
bullet('Doctrine suit quelles migrations ont été appliquées dans la table doctrine_migration_versions')

p()
h(2, '7.2 Créer et exécuter une migration')

code_block("""# 1. Vérifier les différences entre les entités et le schéma actuel
docker compose exec php php bin/console doctrine:schema:validate

# 2. Générer automatiquement la migration
docker compose exec php php bin/console doctrine:migrations:diff

# 3. Appliquer toutes les migrations en attente
docker compose exec php php bin/console doctrine:migrations:migrate --no-interaction

# 4. Voir le statut des migrations
docker compose exec php php bin/console doctrine:migrations:status

# 5. Annuler la dernière migration
docker compose exec php php bin/console doctrine:migrations:migrate prev""")

p()
h(2, '7.3 Analyse de la migration TravelDock')
p('Voici la migration créée pour le projet TravelDock :')

code_block("""<?php
// migrations/Version20240601000001.php
namespace DoctrineMigrations;

use Doctrine\\DBAL\\Schema\\Schema;
use Doctrine\\Migrations\\AbstractMigration;

final class Version20240601000001 extends AbstractMigration
{
    public function getDescription(): string
    {
        return 'Création des tables destination, travel et booking';
    }

    public function up(Schema $schema): void
    {
        // Création de la table destination
        $this->addSql('CREATE TABLE destination (
            id INT AUTO_INCREMENT NOT NULL,
            name VARCHAR(100) NOT NULL,
            country VARCHAR(100) NOT NULL,
            description LONGTEXT DEFAULT NULL,
            iata_code VARCHAR(10) DEFAULT NULL,
            PRIMARY KEY(id)
        ) DEFAULT CHARACTER SET utf8mb4 ENGINE = InnoDB');

        // Création de la table travel avec FK vers destination
        $this->addSql('CREATE TABLE travel (
            id INT AUTO_INCREMENT NOT NULL,
            destination_id INT NOT NULL,           -- Clé étrangère
            title VARCHAR(150) NOT NULL,
            departure_date DATE NOT NULL,
            return_date DATE NOT NULL,
            price NUMERIC(10, 2) NOT NULL,         -- Précision décimale
            max_participants INT NOT NULL,
            description LONGTEXT DEFAULT NULL,
            status VARCHAR(20) NOT NULL DEFAULT \\'draft\\',
            created_at DATETIME NOT NULL,
            INDEX IDX_26B7784B816C6140 (destination_id),   -- Index sur la FK
            PRIMARY KEY(id)
        ) ENGINE = InnoDB');

        // Création de la table booking
        $this->addSql('CREATE TABLE booking (
            id INT AUTO_INCREMENT NOT NULL,
            travel_id INT NOT NULL,
            first_name VARCHAR(100) NOT NULL,
            last_name VARCHAR(100) NOT NULL,
            email VARCHAR(180) NOT NULL,
            phone VARCHAR(20) DEFAULT NULL,
            number_of_persons INT NOT NULL,
            total_price NUMERIC(10, 2) NOT NULL,
            status VARCHAR(20) NOT NULL DEFAULT \\'pending\\',
            booked_at DATETIME NOT NULL,
            PRIMARY KEY(id)
        ) ENGINE = InnoDB');

        // Ajout des contraintes de clés étrangères
        $this->addSql('ALTER TABLE travel
            ADD CONSTRAINT FK_travel_destination
            FOREIGN KEY (destination_id) REFERENCES destination (id)');

        $this->addSql('ALTER TABLE booking
            ADD CONSTRAINT FK_booking_travel
            FOREIGN KEY (travel_id) REFERENCES travel (id)');
    }

    public function down(Schema $schema): void
    {
        // Suppression dans l'ordre inverse (respecter les contraintes FK)
        $this->addSql('ALTER TABLE booking DROP FOREIGN KEY FK_booking_travel');
        $this->addSql('ALTER TABLE travel DROP FOREIGN KEY FK_travel_destination');
        $this->addSql('DROP TABLE booking');
        $this->addSql('DROP TABLE travel');
        $this->addSql('DROP TABLE destination');
    }
}""")

note("La convention de nommage underscore_number_aware de Doctrine convertit automatiquement les propriétés camelCase PHP (departureDate) en colonnes underscore SQL (departure_date).")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 8 — CONTROLLERS ET ROUTES
# ══════════════════════════════════════════════════════════════════════════════

h(1, '8. Les Controllers et les Routes')

h(2, '8.1 Architecture MVC dans Symfony')
p("""Symfony suit le pattern MVC (Model-View-Controller) :
  • Model : les Entités Doctrine (Travel, Destination, Booking)
  • View : les templates Twig (.html.twig)
  • Controller : les classes PHP qui orchestrent la logique

Le Controller reçoit la requête HTTP, interagit avec le Model via les Repositories,
et renvoie une Response avec les données au template View.""")

p()
h(2, '8.2 Le routing par attributs PHP')
p("""Depuis Symfony 6, le routing peut être défini directement dans les classes
Controller via des attributs PHP #[Route(...)], sans fichier YAML séparé :""")

code_block("""use Symfony\\Component\\Routing\\Attribute\\Route;

// Route simple
#[Route('/', name: 'app_home')]
public function index(): Response {}

// Route avec préfixe sur toute la classe
#[Route('/travel', name: 'travel_')]  // → toutes les routes commencent par /travel
class TravelController {}

// Route avec paramètre typé
#[Route('/{id}', name: 'show', requirements: ['id' => '\\\\d+'], methods: ['GET'])]
public function show(Travel $travel): Response {}
//                  ↑ ParamConverter : Symfony charge automatiquement l'entité par son ID

// Route multi-méthodes
#[Route('/new', name: 'new', methods: ['GET', 'POST'])]
public function new(Request $request): Response {}""")

note("requirements: ['id' => '\\\\d+'] garantit que l'URL /travel/abc ne matche pas cette route (seulement les entiers). C'est une sécurité importante.")

p()
h(2, '8.3 Le HomeController')
p('Le HomeController gère la page d\'accueil et illustre l\'injection de services :')

code_block("""<?php
namespace App\\Controller;

use App\\Repository\\DestinationRepository;
use App\\Repository\\TravelRepository;
use Symfony\\Bundle\\FrameworkBundle\\Controller\\AbstractController;
use Symfony\\Component\\HttpFoundation\\Response;
use Symfony\\Component\\Routing\\Attribute\\Route;

class HomeController extends AbstractController
{
    // Autowiring : Symfony injecte automatiquement les repositories
    #[Route('/', name: 'app_home')]
    public function index(
        TravelRepository $travelRepository,
        DestinationRepository $destinationRepository
    ): Response {
        // Appel de méthodes du repository
        $travels      = $travelRepository->findPublished();
        $destinations = $destinationRepository->findWithActiveTravel();

        // render() retourne un objet Response avec le HTML généré
        return $this->render('home/index.html.twig', [
            'travels'      => $travels,       // Variables disponibles dans Twig
            'destinations' => $destinations,
        ]);
    }
}""")

p()
h(2, '8.4 Le TravelController (CRUD complet)')
p("""Le TravelController implémente les opérations CRUD :
Create (new), Read (index/show), Update (edit), Delete (delete).""")

code_block("""<?php
#[Route('/travel', name: 'travel_')]
class TravelController extends AbstractController
{
    // Injection dans le constructeur (une seule fois, pas à chaque méthode)
    public function __construct(
        private readonly EntityManagerInterface $entityManager,
        private readonly TravelRepository $travelRepository,
    ) {}

    // READ ALL — GET /travel
    #[Route('', name: 'index', methods: ['GET'])]
    public function index(): Response
    {
        $travels = $this->travelRepository->findPublished();
        return $this->render('travel/index.html.twig', ['travels' => $travels]);
    }

    // READ ONE — GET /travel/{id}
    // ParamConverter : Symfony cherche Travel::find($id) automatiquement
    // et lève une 404 si non trouvé
    #[Route('/{id}', name: 'show', requirements: ['id' => '\\\\d+'], methods: ['GET'])]
    public function show(Travel $travel): Response
    {
        return $this->render('travel/show.html.twig', ['travel' => $travel]);
    }

    // CREATE — GET /travel/new (formulaire) + POST /travel/new (traitement)
    #[Route('/new', name: 'new', methods: ['GET', 'POST'])]
    public function new(Request $request, DestinationRepository $destRepo): Response
    {
        $travel = new Travel();

        if ($request->isMethod('POST')) {
            // Récupération sécurisée des données POST
            $data = $request->request->all();
            $travel->setTitle($data['title'] ?? '');
            $travel->setPrice($data['price'] ?? '0');
            // ... autres champs ...

            $this->entityManager->persist($travel);
            $this->entityManager->flush();

            // Flash message de succès
            $this->addFlash('success', 'Voyage créé !');
            // Redirection post-POST (pattern PRG)
            return $this->redirectToRoute('travel_show', ['id' => $travel->getId()]);
        }

        return $this->render('travel/new.html.twig', ['travel' => $travel]);
    }

    // DELETE — POST /travel/{id}/delete
    #[Route('/{id}/delete', name: 'delete', requirements: ['id' => '\\\\d+'], methods: ['POST'])]
    public function delete(Request $request, Travel $travel): Response
    {
        // Vérification du token CSRF pour sécuriser la suppression
        if ($this->isCsrfTokenValid('delete' . $travel->getId(),
                                    $request->request->get('_token'))) {
            $this->entityManager->remove($travel);
            $this->entityManager->flush();
            $this->addFlash('success', 'Voyage supprimé.');
        }
        return $this->redirectToRoute('travel_index');
    }
}""")

note("Le pattern PRG (Post-Redirect-Get) est essentiel : après un POST réussi, on redirige toujours vers un GET. Cela évite la soumission multiple du formulaire lors d'un F5.")

p()
h(2, '8.5 Tableau récapitulatif des routes')

table5 = doc.add_table(rows=1, cols=4)
table5.style = 'Table Grid'
for i, txt in enumerate(['Route', 'Méthode', 'Nom', 'Description']):
    table5.rows[0].cells[i].text = txt
    for run in table5.rows[0].cells[i].paragraphs[0].runs:
        run.font.bold = True

routes = [
    ('/',                          'GET',         'app_home',           'Page d\'accueil'),
    ('/travel',                    'GET',         'travel_index',       'Liste des voyages'),
    ('/travel/new',                'GET',         'travel_new',         'Formulaire de création'),
    ('/travel/new',                'POST',        'travel_new',         'Traitement du formulaire'),
    ('/travel/{id}',               'GET',         'travel_show',        'Détail d\'un voyage'),
    ('/travel/{id}/edit',          'GET/POST',    'travel_edit',        'Modification d\'un voyage'),
    ('/travel/{id}/delete',        'POST',        'travel_delete',      'Suppression'),
    ('/destination',               'GET',         'destination_index',  'Liste des destinations'),
    ('/destination/new',           'GET/POST',    'destination_new',    'Nouvelle destination'),
    ('/destination/{id}',          'GET',         'destination_show',   'Détail destination'),
    ('/booking/travel/{id}',       'GET/POST',    'booking_new',        'Réserver un voyage'),
    ('/booking',                   'GET',         'booking_index',      'Liste réservations'),
    ('/booking/{id}',              'GET',         'booking_show',       'Détail réservation'),
    ('/booking/{id}/confirm',      'POST',        'booking_confirm',    'Confirmer'),
    ('/booking/{id}/cancel',       'POST',        'booking_cancel',     'Annuler'),
]
for r in routes:
    row = table5.add_row().cells
    for i, v in enumerate(r):
        row[i].text = v

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 9 — TEMPLATES TWIG
# ══════════════════════════════════════════════════════════════════════════════

h(1, '9. Les Templates Twig')

h(2, '9.1 Le moteur de templates Twig')
p("""Twig est le moteur de templates officiel de Symfony. Il sépare la logique métier
(PHP) de la présentation (HTML), ce qui améliore la lisibilité et la maintenabilité.

Twig utilise 3 types de balises :""")

code_block("""{{ variable }}           {# Affichage d'une variable #}
{% if condition %}       {# Instructions (if, for, block...) #}
{# Ceci est un commentaire #}

{# Filtres : transforment les valeurs #}
{{ travel.price|number_format(2, ',', ' ') }}   {# 1 299,00 #}
{{ travel.createdAt|date('d/m/Y') }}             {# 01/06/2024 #}
{{ travel.description|slice(0, 100) ~ '...' }}   {# Truncate #}

{# Fonctions Twig #}
{{ path('travel_show', {id: travel.id}) }}       {# Génère une URL #}
{{ csrf_token('delete' ~ travel.id) }}           {# Token CSRF #}""")

p()
h(2, '9.2 Le template de base (base.html.twig)')
p("""Le template de base définit la structure HTML commune à toutes les pages.
Il utilise des blocs (block) que les templates enfants peuvent surcharger :""")

code_block("""<!DOCTYPE html>
<html lang="fr">
<head>
    <title>{% block title %}TravelDock{% endblock %}</title>
    {# Bootstrap CSS #}
    <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css">
    {% block stylesheets %}{% endblock %}
</head>
<body>
    <nav class="navbar navbar-dark bg-primary">
        <a class="navbar-brand" href="{{ path('app_home') }}">TravelDock</a>
        {# Navigation avec path() pour les URLs #}
        <a href="{{ path('travel_index') }}">Voyages</a>
    </nav>

    <div class="container mt-4">
        {# Affichage des messages flash (succès, erreurs) #}
        {% for label, messages in app.flashes %}
            {% for message in messages %}
                <div class="alert alert-{{ label }}">{{ message }}</div>
            {% endfor %}
        {% endfor %}

        {% block body %}{% endblock %}   {# Contenu injecté par les enfants #}
    </div>

    {% block javascripts %}{% endblock %}
</body>
</html>""")

p()
h(2, '9.3 Héritage de templates')
p('Les templates enfants étendent le template de base et remplissent les blocs :')

code_block("""{# templates/travel/show.html.twig #}
{% extends 'base.html.twig' %}   {# Héritage #}

{% block title %}{{ travel.title }} - TravelDock{% endblock %}

{% block body %}
<h1>{{ travel.title }}</h1>
<p>Destination : {{ travel.destination }}</p>  {# Appel de __toString() #}

{# Boucle sur les réservations #}
{% for booking in travel.bookings %}
    <p>{{ booking.fullName }} — {{ booking.numberOfPersons }} pers.</p>
{% else %}
    <p>Aucune réservation.</p>
{% endfor %}

{# Condition et chemin relatif #}
{% if travel.availableSeats > 0 %}
    <a href="{{ path('booking_new', {id: travel.id}) }}" class="btn btn-primary">
        Réserver
    </a>
{% else %}
    <button disabled>Complet</button>
{% endif %}

{# Formulaire de suppression avec CSRF #}
<form action="{{ path('travel_delete', {id: travel.id}) }}" method="post">
    <input type="hidden" name="_token" value="{{ csrf_token('delete' ~ travel.id) }}">
    <button type="submit" class="btn btn-danger">Supprimer</button>
</form>
{% endblock %}""")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 10 — MISE EN PRATIQUE
# ══════════════════════════════════════════════════════════════════════════════

h(1, '10. Mise en pratique — Exercices')

h(2, 'Exercice 1 — Démarrage du projet')
p('Suivez ces étapes pour mettre en route TravelDock sur votre machine :')

numbered('Installer Docker Desktop (https://www.docker.com/products/docker-desktop) et le démarrer')
numbered('Ouvrir un terminal dans le dossier symfony-docker/')
numbered('Construire et démarrer les containers :')
code_block('docker compose up -d --build')
numbered('Installer les dépendances Symfony :')
code_block('docker compose exec php composer install')
numbered('Exécuter les migrations :')
code_block('docker compose exec php php bin/console doctrine:migrations:migrate --no-interaction')
numbered('Ouvrir http://localhost:8080 dans le navigateur')

p()
h(2, 'Exercice 2 — Création d\'une destination via phpMyAdmin')
numbered('Ouvrir phpMyAdmin sur http://localhost:8081')
numbered('Naviguer dans la base travelDock → table destination')
numbered('Insérer une ligne : name="Paris", country="France", iata_code="CDG"')
numbered('Vérifier que la destination apparaît sur http://localhost:8080/destination')

p()
h(2, 'Exercice 3 — Création d\'un voyage via l\'interface')
numbered('Sur http://localhost:8080/destination/new, créer la destination "Tokyo / Japon"')
numbered('Sur http://localhost:8080/travel/new, créer un voyage :')
bullet('Titre : "Découverte du Japon"')
bullet('Destination : Tokyo (Japon)')
bullet('Dates : du 15/09/2025 au 29/09/2025')
bullet('Prix : 2 990 €, max 15 participants')
bullet('Statut : Publié')
numbered('Faire une réservation pour ce voyage via http://localhost:8080/travel/{id}')

p()
h(2, 'Exercice 4 — Ajout d\'une méthode de recherche')
p('Ajoutez une méthode dans TravelRepository pour filtrer par budget maximum :')

code_block("""// src/Repository/TravelRepository.php
public function findByMaxPrice(float $maxPrice): array
{
    return $this->createQueryBuilder('t')
        ->andWhere('t.price <= :maxPrice')
        ->setParameter('maxPrice', $maxPrice)
        ->andWhere('t.status = :status')
        ->setParameter('status', 'published')
        ->orderBy('t.price', 'ASC')
        ->getQuery()
        ->getResult();
}""")

p('Puis ajoutez un filtre dans TravelController::index() pour utiliser cette méthode.')

p()
h(2, 'Exercice 5 — Nouvelle entité Review (avancé)')
p('Créez une entité Review permettant aux voyageurs de noter un voyage :')
bullet('Propriétés : id, rating (1-5), comment, createdAt')
bullet('Relation ManyToOne avec Travel')
bullet('Contrainte : rating entre 1 et 5 (Assert\\Range)')
bullet('Générez la migration correspondante')
bullet('Ajoutez un ReviewController avec les routes /review/travel/{id} et /review/{id}')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 11 — RÉFÉRENCE DES COMMANDES
# ══════════════════════════════════════════════════════════════════════════════

h(1, '11. Référence des commandes')

h(2, 'Docker Compose')
table_cmd1 = doc.add_table(rows=1, cols=2)
table_cmd1.style = 'Table Grid'
for i, txt in enumerate(['Commande', 'Description']):
    table_cmd1.rows[0].cells[i].text = txt
    for run in table_cmd1.rows[0].cells[i].paragraphs[0].runs:
        run.font.bold = True
for cmd, desc in [
    ('docker compose up -d', 'Démarrer les containers'),
    ('docker compose up -d --build', 'Reconstruire les images et démarrer'),
    ('docker compose down', 'Arrêter et supprimer les containers'),
    ('docker compose down -v', 'Arrêter + supprimer les volumes (données)'),
    ('docker compose ps', 'Lister les containers'),
    ('docker compose logs -f [service]', 'Afficher les logs'),
    ('docker compose exec php bash', 'Terminal dans le container PHP'),
    ('docker compose restart', 'Redémarrer tous les containers'),
]:
    row = table_cmd1.add_row().cells
    row[0].text = cmd
    row[1].text = desc

p()
h(2, 'Composer (depuis le container PHP)')
table_cmd2 = doc.add_table(rows=1, cols=2)
table_cmd2.style = 'Table Grid'
for i, txt in enumerate(['Commande', 'Description']):
    table_cmd2.rows[0].cells[i].text = txt
    for run in table_cmd2.rows[0].cells[i].paragraphs[0].runs:
        run.font.bold = True
for cmd, desc in [
    ('composer install', 'Installer les dépendances'),
    ('composer update', 'Mettre à jour les dépendances'),
    ('composer require symfony/form', 'Ajouter un package'),
    ('composer require --dev symfony/maker-bundle', 'Ajouter une dépendance de dev'),
]:
    row = table_cmd2.add_row().cells
    row[0].text = cmd
    row[1].text = desc

p()
h(2, 'Console Symfony (php bin/console…)')
table_cmd3 = doc.add_table(rows=1, cols=2)
table_cmd3.style = 'Table Grid'
for i, txt in enumerate(['Commande', 'Description']):
    table_cmd3.rows[0].cells[i].text = txt
    for run in table_cmd3.rows[0].cells[i].paragraphs[0].runs:
        run.font.bold = True
for cmd, desc in [
    ('cache:clear', 'Vider le cache'),
    ('debug:router', 'Lister toutes les routes'),
    ('debug:container', 'Lister les services'),
    ('doctrine:schema:validate', 'Valider le schéma Doctrine'),
    ('doctrine:migrations:migrate', 'Exécuter les migrations'),
    ('doctrine:migrations:diff', 'Générer une migration depuis les entités'),
    ('doctrine:migrations:status', 'Statut des migrations'),
    ('make:entity', 'Générer une entité (avec maker-bundle)'),
    ('make:controller', 'Générer un controller'),
    ('make:migration', 'Générer une migration vide'),
]:
    row = table_cmd3.add_row().cells
    row[0].text = 'php bin/console ' + cmd
    row[1].text = desc

p()
p()
para_fin = doc.add_paragraph()
para_fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_fin = para_fin.add_run('— Fin du support de cours TravelDock —')
run_fin.font.size = Pt(14)
run_fin.font.italic = True
run_fin.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

# ── Sauvegarde ───────────────────────────────────────────────────────────────
output_path = 'c:/Users/Jérémie/symfony-docker/TravelDock_Support_Cours.docx'
doc.save(output_path)
print(f"Document généré : {output_path}")
